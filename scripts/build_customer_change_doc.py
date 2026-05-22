from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "本轮玩法逻辑修改说明-客户版.md"
OUTPUT = ROOT / "docs" / "本轮玩法逻辑修改说明-客户版.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E2EC"
BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"


def set_run_font(run, *, size=None, color=None, bold=None, italic=None, font=BODY_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size, color=None, bold=None):
    style.font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def set_paragraph_spacing(paragraph, *, before=0, after=6, line_spacing=1.10):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def paragraph_border_bottom(paragraph, color="2E74B5", size="12", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGIN_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def next_numbering_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    existing = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId")) and num.get(qn("w:numId")).isdigit()
    ]
    return max(existing, default=0) + 1


def create_numbering_instance(doc: Document, abstract_num_id: str = "7") -> int:
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    num_id = next_numbering_id(doc)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int = 0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def configure_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    set_style_font(styles["Heading 1"], size=16, color=BLUE, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 2"], size=13, color=BLUE, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 3"], size=12, color=DARK_BLUE, bold=True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        if style_name in styles:
            set_style_font(styles[style_name], size=11)
            styles[style_name].paragraph_format.left_indent = Inches(0.5)
            styles[style_name].paragraph_format.first_line_indent = Inches(-0.25)
            styles[style_name].paragraph_format.space_after = Pt(5)
            styles[style_name].paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, after=0)
    run = header.add_run("明日之问｜玩法逻辑修改说明")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer, after=0)
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(footer)
    run = footer.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def split_inline_markup(text: str):
    parts = []
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            parts.append(("plain", text[pos:match.start()]))
        token = match.group(0)
        if token.startswith("`"):
            parts.append(("code", token[1:-1]))
        else:
            parts.append(("bold", token[2:-2]))
        pos = match.end()
    if pos < len(text):
        parts.append(("plain", text[pos:]))
    return parts


def add_marked_text(paragraph, text: str, *, size=11, color=None):
    for kind, chunk in split_inline_markup(text):
        run = paragraph.add_run(chunk)
        if kind == "bold":
            set_run_font(run, size=size, color=color, bold=True)
        elif kind == "code":
            set_run_font(run, size=size, color=color, font="Consolas")
        else:
            set_run_font(run, size=size, color=color)


def is_table_start(lines, idx):
    def is_separator(line: str) -> bool:
        stripped = line.strip()
        if not stripped.startswith("|"):
            return False
        parts = parse_table_row(stripped)
        return bool(parts) and all(re.fullmatch(r":?-{3,}:?", part.strip()) for part in parts)

    return (
        idx + 1 < len(lines)
        and lines[idx].lstrip().startswith("|")
        and lines[idx + 1].lstrip().startswith("|")
        and is_separator(lines[idx + 1])
    )


def parse_table_row(line: str):
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip().replace("<br>", "\n") for cell in stripped.split("|")]


def table_widths(headers, rows):
    col_count = len(headers)
    if col_count == 2:
        return [2600, CONTENT_WIDTH_DXA - 2600]
    if col_count == 3:
        return [2100, 2200, CONTENT_WIDTH_DXA - 4300]
    if col_count == 4:
        return [1800, 1600, 3000, CONTENT_WIDTH_DXA - 6400]
    if col_count == 5:
        return [1900, 1400, 1600, 1600, CONTENT_WIDTH_DXA - 6500]
    if col_count == 6:
        return [1900, 1400, 1400, 1400, 1400, CONTENT_WIDTH_DXA - 7500]
    base = CONTENT_WIDTH_DXA // max(1, col_count)
    widths = [base] * col_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    widths = table_widths(headers, rows)

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        shade_cell(header_cells[i], LIGHT_FILL)
        p = header_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line_spacing=1.10)
        add_marked_text(p, header, size=10.5)
        for run in p.runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if re.fullmatch(r"[0-9./: +\-]+", value) else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, after=0, line_spacing=1.10)
            add_marked_text(p, value, size=10)

    set_table_geometry(table, widths)
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=4)


def add_masthead(doc: Document, title: str, metadata_lines):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=14, after=4)
    run = p.add_run("客户说明文档")
    set_run_font(run, size=10, color=MUTED, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6)
    run = p.add_run(title.replace("# ", ""))
    set_run_font(run, size=23, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=14, line_spacing=1.20)
    run = p.add_run("面向客户说明本轮需求已落地后的玩法结果和数值口径")
    set_run_font(run, size=12.5, color=MUTED)

    for line in metadata_lines:
        clean = line.lstrip("> ").strip()
        if not clean:
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2)
        if "：" in clean:
            label, value = clean.split("：", 1)
            r = p.add_run(f"{label}：")
            set_run_font(r, size=10.5, bold=True)
            r = p.add_run(value.strip())
            set_run_font(r, size=10.5)
        else:
            add_marked_text(p, clean, size=10.5)

    rule = doc.add_paragraph()
    set_paragraph_spacing(rule, before=10, after=14)
    paragraph_border_bottom(rule)


def add_callout(doc: Document, text_lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, color="E2E8F0")
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    set_cell_margins(cell, {"top": 120, "bottom": 120, "start": 160, "end": 160})
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line_spacing=1.15)
    add_marked_text(p, " ".join(text_lines), size=10.5, color=RGBColor(70, 70, 70))
    set_table_geometry(table, [CONTENT_WIDTH_DXA])


def build_doc():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_document(doc)

    title = next((line for line in lines if line.startswith("# ")), "# 明日之问｜本轮玩法逻辑修改说明（客户版）")
    metadata = []
    start_idx = 0
    for i, line in enumerate(lines):
        if line.startswith("> "):
            metadata.append(line)
        if line.startswith("## "):
            start_idx = i
            break
    add_masthead(doc, title, metadata)

    idx = start_idx
    in_code = False
    code_lines = []
    active_numbering_id = None
    last_was_numbered = False
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            last_was_numbered = False
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=2, after=8, line_spacing=1.10)
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, size=10, font="Consolas", color=RGBColor(48, 48, 48))
                in_code = False
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if not stripped:
            last_was_numbered = False
            idx += 1
            continue

        if is_table_start(lines, idx):
            last_was_numbered = False
            headers = parse_table_row(lines[idx])
            rows = []
            idx += 2
            while idx < len(lines) and lines[idx].lstrip().startswith("|"):
                rows.append(parse_table_row(lines[idx]))
                idx += 1
            add_table(doc, headers, rows)
            continue

        if stripped.startswith("## "):
            last_was_numbered = False
            p = doc.add_paragraph(style="Heading 1")
            add_marked_text(p, stripped[3:], size=16, color=BLUE)
            idx += 1
            continue

        if stripped.startswith("### "):
            last_was_numbered = False
            p = doc.add_paragraph(style="Heading 2")
            add_marked_text(p, stripped[4:], size=13, color=BLUE)
            idx += 1
            continue

        if stripped.startswith("- "):
            last_was_numbered = False
            p = doc.add_paragraph(style="List Bullet")
            add_marked_text(p, stripped[2:])
            idx += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            if not last_was_numbered or active_numbering_id is None:
                active_numbering_id = create_numbering_instance(doc)
            p = doc.add_paragraph()
            apply_numbering(p, active_numbering_id)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            set_paragraph_spacing(p, after=5, line_spacing=1.12)
            add_marked_text(p, numbered.group(2))
            last_was_numbered = True
            idx += 1
            continue

        if stripped.startswith("> "):
            last_was_numbered = False
            callout_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("> "):
                callout_lines.append(lines[idx].strip().lstrip("> ").strip())
                idx += 1
            add_callout(doc, callout_lines)
            continue

        last_was_numbered = False
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=6, line_spacing=1.10)
        add_marked_text(p, stripped)
        idx += 1

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_doc()
    print(path)
